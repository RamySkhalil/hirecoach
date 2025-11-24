"""
CV Export Service - Export CVs to various formats (PDF, DOCX, TXT).
"""
from io import BytesIO
from typing import List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class CVExportService:
    """
    Service for exporting CVs to different formats.
    """
    
    @staticmethod
    def get_export_formats() -> List[dict]:
        """Get list of available export formats."""
        formats = [
            {"format": "txt", "name": "Plain Text", "available": True}
        ]
        
        if PDF_AVAILABLE:
            formats.append({"format": "pdf", "name": "PDF Document", "available": True})
        else:
            formats.append({"format": "pdf", "name": "PDF Document (Install reportlab)", "available": False})
        
        if DOCX_AVAILABLE:
            formats.append({"format": "docx", "name": "Word Document", "available": True})
        else:
            formats.append({"format": "docx", "name": "Word Document (Install python-docx)", "available": False})
        
        return formats
    
    @staticmethod
    def export_to_pdf(cv_text: str) -> BytesIO:
        """
        Export CV to PDF format.
        """
        if not PDF_AVAILABLE:
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              rightMargin=0.75*inch, leftMargin=0.75*inch,
                              topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        # Container for elements
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='CustomBody',
                                 parent=styles['Normal'],
                                 fontSize=11,
                                 leading=14,
                                 spaceAfter=6))
        
        styles.add(ParagraphStyle(name='CustomHeading',
                                 parent=styles['Heading1'],
                                 fontSize=16,
                                 leading=20,
                                 spaceAfter=12,
                                 textColor=RGBColor(0, 0, 100)))
        
        # Parse CV text and create PDF
        lines = cv_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 0.1*inch))
                continue
            
            # Check if line is a heading (ALL CAPS or starts with #)
            if line.isupper() and len(line) < 50:
                para = Paragraph(line, styles['CustomHeading'])
            elif line.startswith('#'):
                heading_text = line.lstrip('#').strip()
                para = Paragraph(heading_text, styles['CustomHeading'])
            elif line.startswith('•') or line.startswith('-'):
                para = Paragraph(line, styles['CustomBody'])
            else:
                para = Paragraph(line, styles['CustomBody'])
            
            elements.append(para)
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def export_to_docx(cv_text: str) -> BytesIO:
        """
        Export CV to DOCX format.
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Parse CV text and create DOCX
        lines = cv_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()  # Empty paragraph for spacing
                continue
            
            # Check if line is a heading
            if line.isupper() and len(line) < 50:
                heading = doc.add_heading(line, level=1)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
            elif line.startswith('#'):
                heading_text = line.lstrip('#').strip()
                heading = doc.add_heading(heading_text, level=1)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
            else:
                para = doc.add_paragraph(line)
                para.paragraph_format.space_after = Pt(6)
                
                # Format bullet points
                if line.startswith('•') or line.startswith('-'):
                    para.paragraph_format.left_indent = Inches(0.25)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def export_to_markdown(cv_text: str) -> str:
        """
        Export CV to Markdown format (already in markdown usually).
        """
        return cv_text
